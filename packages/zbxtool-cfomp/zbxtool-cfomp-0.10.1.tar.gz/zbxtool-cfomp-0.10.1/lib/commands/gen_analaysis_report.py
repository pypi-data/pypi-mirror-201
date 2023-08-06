""" 生成zabbix 主机组资源使用率报表
    """
import sys
import time
import argparse
import logging
import math
import traceback
import convertapi
from bisect import bisect_left
from collections import defaultdict, Counter
from itertools import groupby
from operator import itemgetter
from datetime import datetime
from docx.shared import RGBColor
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Inches
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement


RGB_YELLOW = 255, 165, 0
RGB_ORANGE = 255, 102, 0
RGB_RED = 255, 0, 0


def valid_date(s):
    """ 验证 argparse 传入日期格式
    """
    try:
        datetime.strptime(s, "%Y%m%d")
        return s
    except ValueError as exc:
        msg = "Not a valid date format YYYYMMDD: {0}.".format(s)
        raise argparse.ArgumentTypeError(msg) from exc


parser = argparse.ArgumentParser()
parser.add_argument("--start", required=True,
                    help="The Start Date - format YYYYMMDD", type=valid_date)
parser.add_argument("--end", required=True,
                    help="The End Date(Inclusive) - format YYYYMMDD", type=valid_date)
parser.add_argument("--output", type=str, help="output")
parser.add_argument("--topnum", type=int, default=10)
parser.add_argument("--api-secret", type=str, help="convertapi api_secret")
parser.set_defaults(handler=lambda args: main(args))


def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def convert_unit(size):
    """ 将 bytes 转换为易读的单位
        bisect_left 确定 size 在 sizes 中应插入位置 factor
        size 应除以 1024 ** factor, 加上对应的单位字符
    """
    # 各单位表示字符
    units = "BKMGTPE"
    # 各单位对应数值大小, [1024, 1024**2, 1024**3, ...]
    sizes = [1024 ** i for i in range(1, 8)]
    factor = bisect_left(sizes, size)
    return str(round(size / (1024 ** factor), 2)) + units[factor]


def add_summarize_table(doc, style):
    """
    创建"主机组资源利用率概览"表。
    为实现分页添加表头功能, 将创建表的动作独立出来, 并自动添加表头。
    Return: table
    """
    table = doc.add_table(
        rows=1, cols=12, style=style)
    table.autofit = False
    table.cell(0, 0).width, table.cell(0, 0).text = Inches(1), "主机组名称"
    table.cell(0, 1).width, table.cell(0, 1).text = Inches(0.5), "主机数量"
    table.cell(0, 2).width, table.cell(0, 2).text = Inches(0.75), "CPU平均利用率"
    table.cell(0, 3).width, table.cell(0, 3).text = Inches(0.75), "内存总量"
    table.cell(0, 4).width, table.cell(0, 4).text = Inches(0.75), "内存最高利用率"
    table.cell(0, 5).width, table.cell(0, 5).text = Inches(0.75), "内存最低利用率"
    table.cell(0, 6).width, table.cell(0, 6).text = Inches(0.75), "内存平均利用率"
    table.cell(0, 7).width, table.cell(0, 7).text = Inches(0.75), "磁盘总量"
    table.cell(0, 8).width, table.cell(0, 8).text = Inches(0.75), "磁盘最高使用率"
    table.cell(0, 9).width, table.cell(0, 9).text = Inches(0.75), "磁盘最低使用率"
    table.cell(0, 10).width, table.cell(0, 10).text = Inches(0.75), "磁盘平均使用率"
    table.cell(0, 11).width, table.cell(0, 11).text = Inches(0.75), "严重告警数量"
    return table


def add_vm_table(doc, style):
    """ 创建cpu使用率低于1%的虚拟机表格
        为实现分页添加表头功能, 将创建表的动作独立出来, 并自动添加表头。
    """
    table = doc.add_table(
        rows=1, cols=3, style=style)
    table.cell(0, 0).width, table.cell(0, 0).text = Inches(3.0), "主机组名称"
    table.cell(0, 1).width, table.cell(0, 1).text = Inches(4.0), "主机名称"
    table.cell(0, 2).width, table.cell(0, 2).text = Inches(0.7), "CPU平均使用率"
    return table


def colored_cell(cell, value):
    """
        根据 value 设置cell的字体颜色。
        85 < value <= 90, RGB_YELLOW
        90 < value <= 95, RGB_ORANGE
        95 < value, RGB_RED
    """
    if value >= 100:
        value = 99.99
    run = cell.paragraphs[0].add_run(str("%.2f" % value)+"%")
    if 85 < value <= 90:
        run.font.color.rgb = RGBColor(*RGB_YELLOW)
    elif 90 < value <= 95:
        run.font.color.rgb = RGBColor(*RGB_ORANGE)
    elif 95 < value:
        run.font.color.rgb = RGBColor(*RGB_RED)


def get_word(api, server_hostid, path, start, end, topnum):
    """" 生成word统计报表 """

    end_timestamp = time.mktime(time.strptime(end, "%Y%m%d"))
    start_timestamp = time.mktime(time.strptime(start, "%Y%m%d"))
    document = Document()

    # 设置正文中的字体 - 微软雅黑
    document.styles["Normal"].font.name = "微软雅黑"
    document.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    paragraph_cover = document.add_paragraph("")
    paragraph_cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cover = paragraph_cover.add_run("\n鑫运运管平台\n监控统计分析月报\n\n")
    run_cover.bold = True
    run_cover.font.size = Pt(36)
    run_cover.font.color.rgb = RGBColor(79, 129, 189)
    run_time = paragraph_cover.add_run(
        "\n\n" + end[0:4] + "年" + end[4:6] + "月")
    run_time.bold = True
    run_time.font.size = Pt(18)
    run_time.font.color.rgb = RGBColor(79, 129, 189)

    document.add_page_break()
    # 1.汇总信息页
    run_1 = document.add_heading("", level=1).add_run("一、汇总信息")
    run_1.font.name = "微软雅黑"
    run_1.font.size = Pt(20)
    run_1._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 1.1表头
    table_total = document.add_table(
        rows=4, cols=2, style="Medium Shading 1 Accent 1")
    table_total.cell(0, 0).text = "统计日期"
    table_total.cell(1, 0).text = "主机组数量"
    table_total.cell(2, 0).text = "主机数量"
    table_total.cell(3, 0).text = "严重告警数量"

    # 1.2表内数据
    table_total.cell(0, 1).text = "{} - {}".format(
        time.strftime("%Y/%m/%d", time.strptime(start, "%Y%m%d")),
        time.strftime("%Y/%m/%d", time.strptime(end, "%Y%m%d")))

    # 获取主机组
    host_groups = api.hostgroup.get({
        "selectHosts": ["hostid", "name"],
        "real_hosts": True,
        "with_monitored_items": True,
        "filter": {"flags": 0}
    })

    # 主机组总数量
    groups_num = len(host_groups)
    # 主机总数量
    hosts_sum = []
    for grp in host_groups:
        hosts_sum += [host["hostid"] for host in grp["hosts"]]
        hosts_sum_num = len(set(hosts_sum))
    # 获取严重告警数量
    event_sum_num = api.event.get({
        "countOutput": True,
        "value": 1,
        "severities": [3, 4, 5],
        "time_from": start_timestamp,
        "time_till": end_timestamp
    })

    table_total.cell(1, 1).text = str(groups_num)
    table_total.cell(2, 1).text = str(hosts_sum_num)
    table_total.cell(3, 1).text = str(event_sum_num)

    run_event_number = document.add_paragraph("")
    # run_event_number.paragraph_format.line_spacing = 1.0
    run_event_number.paragraph_format.space_before = 15
    # event_detail_number.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # run_detail_number = run_event_number.add_run("\n各级别告警数量")
    # run_detail_number.font.name = "微软雅黑"
    # run_detail_number.font.size = Pt(14)
    # run_detail_number._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    table_detail_number = document.add_table(
        rows=3, cols=6, style="Medium Shading 1 Accent 1")
    cell_left = table_detail_number.cell(0, 0)
    cell_right = table_detail_number.cell(0, 5)
    table_detail_header = cell_left.merge(cell_right)
    table_detail_header.text = "各级别告警数量"
    table_detail_number.cell(1, 0).text = "未分类"
    table_detail_number.cell(1, 1).text = "通知"
    table_detail_number.cell(1, 2).text = "警示"
    table_detail_number.cell(1, 3).text = "严重"
    table_detail_number.cell(1, 4).text = "危险"
    table_detail_number.cell(1, 5).text = "灾难"

    # 获取对应告警级别数量
    for severity in range(6):
        event_num = api.event.get({
            "countOutput": True,
            "value": 1,
            "severities": [severity],
            "time_from": start_timestamp,
            "time_till": end_timestamp
        })
        table_detail_number.cell(2, severity).text = str(event_num)

    event_note = document.add_paragraph("")
    run_cover = event_note.add_run("注: `严重`、`危险`、`灾难` 三个等级的告警纳入严重告警统计")
    run_cover.bold = True
    run_cover.font.size = Pt(10)
    run_cover.font.color.rgb = RGBColor(0, 139, 0)

    # 严重告警数量表格
    document.add_page_break()
    run_event = document.add_heading("", level=1).add_run("二、严重告警数量排行")
    run_event.font.name = "微软雅黑"
    run_event.font.size = Pt(20)
    run_event._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 新增2级标题
    run_event_1_desc = document.add_heading("", level=2).add_run(
        "1、严重告警数量最多的{}个主机组".format(topnum))
    run_event_1_desc.font.name = "微软雅黑"
    run_event_1_desc._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run_event_1_desc.font.size = Pt(16)

    # 插入表格
    event_table_desc_group = document.add_table(
        rows=1, cols=3, style="Medium Shading 1 Accent 1")
    event_table_desc_group.cell(0, 0).width = Inches(5)
    event_table_desc_group.cell(0, 1).width = Inches(1)
    event_table_desc_group.cell(0, 2).width = Inches(2.5)
    event_table_desc_group.cell(0, 0).text = "主机组名称"
    event_table_desc_group.cell(0, 1).text = "主机数量"
    event_table_desc_group.cell(0, 2).text = "严重告警数量"

    document.add_page_break()

    # 严重告警数量最多的主机
    logging.info("严重告警数量排行, 主机维度")
    run_event_2_desc = document.add_heading("", level=2).add_run(
        "2、严重告警数量最多的{}台主机".format(topnum))
    run_event_2_desc.font.name = "微软雅黑"
    run_event_2_desc._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run_event_2_desc.font.size = Pt(16)
    event_table_desc_host = document.add_table(
        rows=1, cols=3, style="Medium Shading 1 Accent 1")
    event_table_desc_host.cell(0, 0).width = Inches(3.0)
    event_table_desc_host.cell(0, 1).width = Inches(3.4)
    event_table_desc_host.cell(0, 2).width = Inches(1.3)
    event_table_desc_host.cell(0, 0).text = "主机组名称"
    event_table_desc_host.cell(0, 1).text = "主机名称"
    event_table_desc_host.cell(0, 2).text = "严重告警数量"

    # 2.详细统计信息页
    # 2.1 表头
    document.add_page_break()
    run_overview = document.add_heading("", level=1).add_run("三、主机组资源利用率概览")
    run_overview.font.name = "微软雅黑"
    run_overview.font.size = Pt(20)
    run_overview._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 2.3 获取数据
    # 2.3.1 获取 application_id
    apps = api.application.get({
        "hostids": server_hostid,
        "filter": {
            "name": ["Memory aggregation", "Filesystem aggregation", "CPU aggregation"]
        },
        "output": ["applicationid"]
    })

    # 2.3.2 获取 zabbix server 中各 hostgroup的聚合 item
    items = api.item.get({
        "hostids": server_hostid,
        "applicationids": [app["applicationid"] for app in apps],
        "output": ["name", "kay_"],
        "monitored": True,  # 已启用item
        "filter": {
            "state": 0  # 0 - normal, 1 - Not supported
        }
    })

    # 根据 Average cpu utilization 监控项确定主机组
    valid_hostgroup_names = [item["name"].split("group")[1].strip()
                             for item in items
                             if item["name"].startswith("Average cpu utilization in group")]
    host_groups = [g for g in host_groups if g["name"]
                   in valid_hostgroup_names]

    # 按主机数量排序主机组
    host_groups.sort(key=lambda x: len(x["hosts"]), reverse=True)

    # 2.3.3 设置按主机组维度统计数据的变量
    # 主机组维度按内存使用率avg数组
    memory_top_group = []
    # cpu利用率数组（主机组维度）
    cpu_top_group = []
    # 磁盘使用率（主机组维度）
    filesystem_top_group = []
    # 告警数量 (主机组维度)
    event_count_group = []
    # 主机维度按内存使用率avg数组
    memory_top_host = []
    # cpu利用率数组（主机维度）
    cpu_top_host = []
    # 磁盘使用率（主机维度）
    filesystem_top_host = []
    # 告警数量 (主机组维度)
    event_count_host = []

    # 2.3.4 填充表格数据
    summarize_row_count = 0
    for index, group in enumerate(host_groups):
        group_name = group["name"]
        logging.info("正在处理数据……主机组：%s", group_name)
        logging.info("开始时间：%s", str(datetime.now()))
        if summarize_row_count == 0:
            summarize_table = add_summarize_table(
                document, "Medium Shading 1 Accent 1")
        host_num = len(group["hosts"])
        row = summarize_table.add_row()
        row.cells[0].text = group_name
        row.cells[1].text = str(host_num)
        # group_name 5个字一行, 计算共占多少行
        summarize_row_count += math.ceil(len(group_name) / 5)
        # 获取cpu利用率
        item_cpu_name = f"Average cpu utilization in group {group_name}"
        itemid = [item["itemid"]
                  for item in items
                  if item["name"] == item_cpu_name]
        _, _, avg_v = getcalc(
            api, itemid, start_timestamp, end_timestamp)
        colored_cell(row.cells[2], avg_v)

        # 保留信息
        cpu_top_group.append({
            "groupname": group_name, "hostnum": host_num,
            "cpu_utilization": avg_v})

        # 获取内存总量
        item_total_memory_name = f"Total memory in group {group_name}"
        itemid = [item["itemid"]
                  for item in items
                  if item["name"] == item_total_memory_name]
        _, _, avg_v = getcalc(
            api, itemid, start_timestamp, end_timestamp)

        row.cells[3].text = convert_unit(avg_v)
        memory_dic = {"groupname": group_name,
                      "hostnum": host_num, "memory_total": avg_v}
        # 获取内存利用率
        item_utilization_memory_name = f"Memory utilization in group {group_name}"
        itemid = [item["itemid"]
                  for item in items
                  if item["name"] == item_utilization_memory_name]
        min_v, max_v, avg_v = getcalc(
            api, itemid, start_timestamp, end_timestamp)

        colored_cell(row.cells[4], max_v)
        colored_cell(row.cells[5], min_v)
        colored_cell(row.cells[6], avg_v)
        memory_dic["memory_utilization"] = avg_v
        memory_top_group.append(memory_dic)

        # 获取磁盘总量
        item_total_filesystem_name = f"Total disk space in {group_name}"
        itemid = [item["itemid"]
                  for item in items
                  if item["name"] == item_total_filesystem_name]
        _, _, avg_v = getcalc(
            api, itemid, start_timestamp, end_timestamp)

        row.cells[7].text = convert_unit(avg_v)
        filesystem_dic = {"groupname": group_name,
                          "hostnum": host_num, "filesystem_total": avg_v}
        # 获取磁盘使用率
        item_utilization_filesystem_name = f"Used disk space in {group_name} (percentage)"
        itemid = [item["itemid"]
                  for item in items
                  if item["name"] == item_utilization_filesystem_name]
        min_v, max_v, avg_v = getcalc(
            api, itemid, start_timestamp, end_timestamp)

        colored_cell(row.cells[8], max_v)
        colored_cell(row.cells[9], min_v)
        colored_cell(row.cells[10], avg_v)
        filesystem_dic["filesystem_utilization"] = avg_v
        filesystem_top_group.append(filesystem_dic)

        # 按主机维度处理信息，包括过滤警告，以及获取对应主机的分析数据
        host_items = api.item.get({
            "hostids": [host['hostid'] for host in group["hosts"]],
            "filter": {
                "key_": [
                    "vfs.fs.totalsize",
                    "vfs.fs.usedsize",
                    "system.cpu.util[,idle]",
                    "vm.memory.size[used]",
                    "vm.memory.size[total]"
                ],
                "state": 0
            },
            "output": ["name", "key_", "hostid"],
            "monitored": True
        })

        group_host_keys = defaultdict(dict)
        for host_item in host_items:
            host_name = [host["name"] for host in group["hosts"]
                         if host["hostid"] == host_item["hostid"]][0]
            group_host_keys[host_name][host_item["key_"]] = host_item["itemid"]

        for host_name, host_keys in group_host_keys.items():

            # 获取主机分析数据
            # 内存 used 、 total
            if host_keys.get("vm.memory.size[total]"):

                _, _, mem_avg_used = getcalc(api,
                                             host_keys["vm.memory.size[used]"],
                                             start_timestamp,
                                             end_timestamp)

                _, _, mem_avg_total = getcalc(api,
                                              host_keys["vm.memory.size[total]"],
                                              start_timestamp,
                                              end_timestamp)
                if mem_avg_total != 0:
                  # 内存 使用率
                    mem_avg_utilization = 100 * mem_avg_used / mem_avg_total
                    memory_top_host.append({"hostname": host_name,
                                            "memory_utilization": mem_avg_utilization,
                                            "memory_total": mem_avg_total,
                                            "groupname": group_name})
            # cpu 使用率
            if host_keys.get("system.cpu.util[,idle]"):
                _, _, cpu_avg_idle = getcalc(api,
                                             host_keys["system.cpu.util[,idle]"],
                                             start_timestamp,
                                             end_timestamp)
                if cpu_avg_idle != 0:
                    cpu_top_host.append({"hostname": host_name,
                                         "hostid": host_item["hostid"],
                                         "cpu_utilization": 100 - cpu_avg_idle,
                                         "groupname": group_name})
            # 磁盘 used 、 total
            if host_keys.get("vfs.fs.totalsize") and host_keys.get("vfs.fs.usedsize"):
                _, _, disk_avg_used = getcalc(api,
                                              host_keys["vfs.fs.usedsize"],
                                              start_timestamp,
                                              end_timestamp)
                _, _, disk_avg_total = getcalc(api,
                                               host_keys["vfs.fs.totalsize"],
                                               start_timestamp,
                                               end_timestamp)
                # 磁盘 使用率
                if disk_avg_used != 0:
                    disk_avg_utilization = 100 * disk_avg_used / disk_avg_total
                    filesystem_top_host.append({"hostname": host_name,
                                                "filesystem_utilization": disk_avg_utilization,
                                                "filesystem_total": disk_avg_total,
                                                "groupname": group_name})
        events = api.event.get({
            "output": ["eventid"],
            "selectHosts": ["name"],
            "hostids": [host['hostid'] for host in group["hosts"]],
            "value": 1,
            "severities": [3, 4, 5],
            "time_from": start_timestamp,
            "time_till": end_timestamp
        })
        row.cells[11].text = str(len(events))
        # 主机组维度 严重告警事件数量
        event_count_dic = {"groupname": group_name,
                           "hostnum": host_num,
                           "events_count": len(events)}
        event_count_group.append(event_count_dic)
        # 主机维度 严重告警事件数量
        events_by_host = Counter(e['hosts'][0]['name']
                                 for e in events if e['hosts'])
        for host_name in events_by_host:
            event_count_host.append({"hostname": host_name,
                                     "events_count": events_by_host[host_name],
                                     "groupname": group_name})

        if index == len(host_groups) - 1:
            document.add_page_break()
        elif summarize_row_count >= 18:
            summarize_row_count = 0
            document.add_page_break()

    # 主机组按严重告警数量排序desc
    event_count_group.sort(key=lambda x: x["events_count"], reverse=True)
    for i in range(min(topnum, len(event_count_group))):
        row = event_table_desc_group.add_row()
        row.cells[0].text = event_count_group[i]["groupname"]
        row.cells[1].text = str(event_count_group[i]["hostnum"])
        row.cells[2].text = str(event_count_group[i]["events_count"])

    event_count_host.sort(key=lambda x: x["events_count"], reverse=True)
    for i in range(min(topnum, len(event_count_host))):
        row = event_table_desc_host.add_row()
        row.cells[0].text = event_count_host[i]["groupname"]
        row.cells[1].text = event_count_host[i]["hostname"]
        row.cells[2].text = str(event_count_host[i]["events_count"])

    # 3. 内存使用率排行
    run_memory = document.add_heading("", level=1).add_run("四、内存使用率排行")
    run_memory.font.name = "微软雅黑"
    run_memory.font.size = Pt(20)
    run_memory._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    run_memory_group_desc = document.add_heading("", level=2).add_run(
        "1、内存使用率最高的{}个主机组".format(topnum))
    run_memory_group_desc.font.name = "微软雅黑"
    run_memory_group_desc._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run_memory_group_desc.font.size = Pt(16)
    # 插入表格 按内存使用率排序desc
    table_memory_group_desc = document.add_table(
        rows=1, cols=4, style="Medium Shading 1 Accent 1")
    table_memory_group_desc.cell(0, 0).width = Inches(5)
    table_memory_group_desc.cell(0, 1).width = Inches(0.5)
    table_memory_group_desc.cell(0, 2).width = Inches(1.7)
    table_memory_group_desc.cell(0, 3).width = Inches(1.3)
    table_memory_group_desc.cell(0, 0).text = "主机组名称"
    table_memory_group_desc.cell(0, 1).text = "主机数量"
    table_memory_group_desc.cell(0, 2).text = "内存平均使用率"
    table_memory_group_desc.cell(0, 3).text = "内存总量"

    memory_top_group.sort(key=lambda x: x["memory_utilization"], reverse=True)
    for i in range(min(topnum, len(memory_top_group))):
        row = table_memory_group_desc.add_row()
        row.cells[0].text = memory_top_group[i]["groupname"]
        row.cells[1].text = str(memory_top_group[i]["hostnum"])
        colored_cell(row.cells[2], memory_top_group[i]["memory_utilization"])
        row.cells[3].text = str(convert_unit(
            memory_top_group[i]["memory_total"]))

    document.add_page_break()
    run_memory_group_asc = document.add_heading("", level=2).add_run(
        "2、内存使用率最低的{}个主机组".format(topnum))
    run_memory_group_asc.font.name = "微软雅黑"
    run_memory_group_asc._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run_memory_group_asc.font.size = Pt(16)
    # 插入表格 按内存使用率排序asc
    table_memory_group_asc = document.add_table(
        rows=1, cols=4, style="Medium Shading 1 Accent 1")
    table_memory_group_asc.cell(0, 0).width = Inches(5)
    table_memory_group_asc.cell(0, 1).width = Inches(0.5)
    table_memory_group_asc.cell(0, 2).width = Inches(1.7)
    table_memory_group_asc.cell(0, 3).width = Inches(1.3)
    table_memory_group_asc.cell(0, 0).text = "主机组名称"
    table_memory_group_asc.cell(0, 1).text = "主机数量"
    table_memory_group_asc.cell(0, 2).text = "内存平均使用率"
    table_memory_group_asc.cell(0, 3).text = "内存总量"

    memory_top_group.sort(key=lambda x: x["memory_utilization"])
    for i in range(min(topnum, len(memory_top_group))):
        row = table_memory_group_asc.add_row()
        row.cells[0].text = memory_top_group[i]["groupname"]
        row.cells[1].text = str(memory_top_group[i]["hostnum"])
        colored_cell(row.cells[2], memory_top_group[i]["memory_utilization"])
        row.cells[3].text = str(convert_unit(
            memory_top_group[i]["memory_total"]))

    document.add_page_break()
    run_memory_host_desc = document.add_heading("", level=2).add_run(
        "3、内存使用率最高的{}台主机".format(topnum))
    run_memory_host_desc.font.name = "微软雅黑"
    run_memory_host_desc._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run_memory_host_desc.font.size = Pt(16)
    # 插入表格
    table_memory_host_desc = document.add_table(
        rows=1, cols=4, style="Medium Shading 1 Accent 1")
    table_memory_host_desc.cell(0, 0).width = Inches(3.0)
    table_memory_host_desc.cell(0, 1).width = Inches(3.4)
    table_memory_host_desc.cell(0, 2).width = Inches(1.3)
    table_memory_host_desc.cell(0, 3).width = Inches(1.3)
    table_memory_host_desc.cell(0, 0).text = "主机组名称"
    table_memory_host_desc.cell(0, 1).text = "主机名称"
    table_memory_host_desc.cell(0, 2).text = "内存平均使用率"
    table_memory_host_desc.cell(0, 3).text = "内存总量"

    memory_top_host.sort(key=itemgetter("memory_utilization"))
    memory_top_host_groupby = []
    for hostname, hosts_iter in groupby(memory_top_host, key=itemgetter("hostname")):
        hosts = list(hosts_iter)
        memory_top_host_groupby.append({
            "hostname": hostname,
            "memory_utilization": hosts[0]["memory_utilization"],
            "memory_total": hosts[0]["memory_total"],
            "groupname": ','.join(h['groupname'] for h in hosts)
        })
    memory_top_host_groupby.sort(
        key=itemgetter("memory_utilization"), reverse=True)
    for i in range(min(topnum, len(memory_top_host))):
        row = table_memory_host_desc.add_row()
        row.cells[0].text = memory_top_host_groupby[i]["groupname"]
        row.cells[1].text = memory_top_host_groupby[i]["hostname"]
        colored_cell(
            row.cells[2], memory_top_host_groupby[i]["memory_utilization"])
        row.cells[3].text = str(convert_unit(
            memory_top_host_groupby[i]["memory_total"]))

    document.add_page_break()
    run_memory_host_asc = document.add_heading("", level=2).add_run(
        "4、内存使用率最低的{}台主机".format(topnum))
    run_memory_host_asc.font.name = "微软雅黑"
    run_memory_host_asc._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run_memory_host_asc.font.size = Pt(16)
    # 插入表格
    table_memory_host_asc = document.add_table(
        rows=1, cols=4, style="Medium Shading 1 Accent 1")
    table_memory_host_asc.cell(0, 0).width = Inches(3.0)
    table_memory_host_asc.cell(0, 1).width = Inches(3.4)
    table_memory_host_asc.cell(0, 2).width = Inches(1.3)
    table_memory_host_asc.cell(0, 3).width = Inches(1.3)
    table_memory_host_asc.cell(0, 0).text = "主机组名称"
    table_memory_host_asc.cell(0, 1).text = "主机名称"
    table_memory_host_asc.cell(0, 2).text = "内存平均使用率"
    table_memory_host_asc.cell(0, 3).text = "内存总量"

    memory_top_host_groupby.sort(key=itemgetter("memory_utilization"))
    for i in range(min(topnum, len(memory_top_host))):
        row = table_memory_host_asc.add_row()
        row.cells[0].text = memory_top_host_groupby[i]["groupname"]
        row.cells[1].text = memory_top_host_groupby[i]["hostname"]
        colored_cell(
            row.cells[2], memory_top_host_groupby[i]["memory_utilization"])
        row.cells[3].text = str(convert_unit(
            memory_top_host_groupby[i]["memory_total"]))

    document.add_page_break()
    run_cpu = document.add_heading("", level=1).add_run("五、CPU使用率排行")
    run_cpu.font.name = "微软雅黑"
    run_cpu.font.size = Pt(20)
    run_cpu._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run_cpu_group_desc = document.add_heading("", level=2).add_run(
        "1、CPU使用率最高的{}个主机组".format(topnum))
    run_cpu_group_desc.font.name = "微软雅黑"
    run_cpu_group_desc._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run_cpu_group_desc.font.size = Pt(16)
    # 插入表格 按cpu使用率排序主机组维度 desc
    table_cpu_group_desc = document.add_table(
        rows=1, cols=3, style="Medium Shading 1 Accent 1")
    table_cpu_group_desc.cell(0, 0).width = Inches(5)
    table_cpu_group_desc.cell(0, 1).width = Inches(1.0)
    table_cpu_group_desc.cell(0, 2).width = Inches(1.2)
    table_cpu_group_desc.cell(0, 0).text = "主机组名称"
    table_cpu_group_desc.cell(0, 1).text = "主机数量"
    table_cpu_group_desc.cell(0, 2).text = "CPU平均使用率"
    cpu_top_group.sort(key=lambda x: x["cpu_utilization"], reverse=True)
    for i in range(min(topnum, len(cpu_top_group))):
        row = table_cpu_group_desc.add_row()
        row.cells[0].text = cpu_top_group[i]["groupname"]
        row.cells[1].text = str(cpu_top_group[i]["hostnum"])
        colored_cell(row.cells[2], cpu_top_group[i]["cpu_utilization"])

    document.add_page_break()
    run_cpu_group_asc = document.add_heading("", level=2).add_run(
        "2、CPU使用率最低的{}个主机组".format(topnum))
    run_cpu_group_asc.font.name = "微软雅黑"
    run_cpu_group_asc._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run_cpu_group_asc.font.size = Pt(16)
    # 插入表格 按cpu使用率排序 主机组维度 asc
    table_cpu_group_asc = document.add_table(
        rows=1, cols=3, style="Medium Shading 1 Accent 1")
    table_cpu_group_asc.cell(0, 0).width = Inches(5)
    table_cpu_group_asc.cell(0, 1).width = Inches(1.0)
    table_cpu_group_asc.cell(0, 2).width = Inches(1.2)
    table_cpu_group_asc.cell(0, 0).text = "主机组名称"
    table_cpu_group_asc.cell(0, 1).text = "主机数量"
    table_cpu_group_asc.cell(0, 2).text = "CPU平均使用率"

    cpu_top_group.sort(key=lambda x: x["cpu_utilization"])
    for i in range(min(topnum, len(cpu_top_group))):
        row = table_cpu_group_asc.add_row()
        row.cells[0].text = cpu_top_group[i]["groupname"]
        row.cells[1].text = str(cpu_top_group[i]["hostnum"])
        colored_cell(row.cells[2], cpu_top_group[i]["cpu_utilization"])

    document.add_page_break()
    run_cpu_host_desc = document.add_heading("", level=2).add_run(
        "3、CPU使用率最高的{}台主机".format(topnum))
    run_cpu_host_desc.font.name = "微软雅黑"
    run_cpu_host_desc._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run_cpu_host_desc.font.size = Pt(16)

    # 插入表格 cpu使用率主机维度 desc
    table_cpu_host_desc = document.add_table(
        rows=1, cols=3, style="Medium Shading 1 Accent 1")
    table_cpu_host_desc.cell(0, 0).width = Inches(3.5)
    table_cpu_host_desc.cell(0, 1).width = Inches(3.5)
    table_cpu_host_desc.cell(0, 2).width = Inches(0.7)
    table_cpu_host_desc.cell(0, 0).text = "主机组名称"
    table_cpu_host_desc.cell(0, 1).text = "主机名称"
    table_cpu_host_desc.cell(0, 2).text = "CPU平均使用率"

    cpu_top_host.sort(key=itemgetter("cpu_utilization"))
    cpu_top_host_groupby = []
    for hostname, hosts_iter in groupby(cpu_top_host, key=itemgetter("hostname")):
        hosts = list(hosts_iter)
        cpu_top_host_groupby.append({
            "hostname": hostname,
            "cpu_utilization": hosts[0]["cpu_utilization"],
            "groupname": ','.join(h['groupname'] for h in hosts)
        })
    cpu_top_host_groupby.sort(
        key=itemgetter("cpu_utilization"), reverse=True)
    for i in range(min(topnum, len(cpu_top_host_groupby))):
        row = table_cpu_host_desc.add_row()
        row.cells[0].text = cpu_top_host_groupby[i]["groupname"]
        row.cells[1].text = cpu_top_host_groupby[i]["hostname"]
        colored_cell(row.cells[2], cpu_top_host_groupby[i]["cpu_utilization"])

    document.add_page_break()
    run_cpu_host_asc = document.add_heading("", level=2).add_run(
        "4、CPU使用率最低的{}台主机".format(topnum))
    run_cpu_host_asc.font.name = "微软雅黑"
    run_cpu_host_asc._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run_cpu_host_asc.font.size = Pt(16)

    # 插入表格
    table_cpu_host_asc = document.add_table(
        rows=1, cols=3, style="Medium Shading 1 Accent 1")
    table_cpu_host_asc.cell(0, 0).width = Inches(3.5)
    table_cpu_host_asc.cell(0, 1).width = Inches(3.5)
    table_cpu_host_asc.cell(0, 2).width = Inches(0.7)
    table_cpu_host_asc.cell(0, 0).text = "主机组名称"
    table_cpu_host_asc.cell(0, 1).text = "主机名称"
    table_cpu_host_asc.cell(0, 2).text = "CPU平均使用率"

    cpu_top_host_groupby.sort(key=itemgetter("cpu_utilization"))
    for i in range(min(topnum, len(cpu_top_host_groupby))):
        row = table_cpu_host_asc.add_row()
        row.cells[0].text = cpu_top_host_groupby[i]["groupname"]
        row.cells[1].text = cpu_top_host_groupby[i]["hostname"]
        colored_cell(row.cells[2], cpu_top_host_groupby[i]["cpu_utilization"])

    document.add_page_break()
    run_disk = document.add_heading("", level=1).add_run("六、磁盘使用率排行")
    run_disk.font.name = "微软雅黑"
    run_disk.font.size = Pt(20)
    run_disk._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    run_disk_group_desc = document.add_heading("", level=2).add_run(
        "1、磁盘使用率最高的{}个主机组".format(topnum))
    run_disk_group_desc.font.name = "微软雅黑"
    run_disk_group_desc._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run_disk_group_desc.font.size = Pt(16)

    # 插入表格   主机组按磁盘使用率排序desc
    table_disk_group_desc = document.add_table(
        rows=1, cols=4, style="Medium Shading 1 Accent 1")
    table_disk_group_desc.cell(0, 0).width = Inches(5)
    table_disk_group_desc.cell(0, 1).width = Inches(0.5)
    table_disk_group_desc.cell(0, 2).width = Inches(1.7)
    table_disk_group_desc.cell(0, 3).width = Inches(1.3)
    table_disk_group_desc.cell(0, 0).text = "主机组名称"
    table_disk_group_desc.cell(0, 1).text = "主机数量"
    table_disk_group_desc.cell(0, 2).text = "磁盘平均使用率"
    table_disk_group_desc.cell(0, 3).text = "磁盘总量"

    filesystem_top_group.sort(
        key=lambda x: x["filesystem_utilization"], reverse=True)
    for i in range(min(topnum, len(filesystem_top_group))):
        row = table_disk_group_desc.add_row()
        row.cells[0].text = filesystem_top_group[i]["groupname"]
        row.cells[1].text = str(filesystem_top_group[i]["hostnum"])
        colored_cell(row.cells[2], filesystem_top_group[i]
                     ["filesystem_utilization"])
        row.cells[3].text = str(convert_unit(
            filesystem_top_group[i]["filesystem_total"]))

    document.add_page_break()
    run_disk_group_asc = document.add_heading("", level=2).add_run(
        "2、磁盘使用率最低的{}个主机组".format(topnum))
    run_disk_group_asc.font.name = "微软雅黑"
    run_disk_group_asc._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run_disk_group_asc.font.size = Pt(16)

    # 插入表格  主机组按磁盘使用率排序asc
    table_disk_group_asc = document.add_table(
        rows=1, cols=4, style="Medium Shading 1 Accent 1")
    table_disk_group_asc.cell(0, 0).width = Inches(5)
    table_disk_group_asc.cell(0, 1).width = Inches(0.5)
    table_disk_group_asc.cell(0, 2).width = Inches(1.7)
    table_disk_group_asc.cell(0, 3).width = Inches(1.3)
    table_disk_group_asc.cell(0, 0).text = "主机组名称"
    table_disk_group_asc.cell(0, 1).text = "主机数量"
    table_disk_group_asc.cell(0, 2).text = "磁盘平均使用率"
    table_disk_group_asc.cell(0, 3).text = "磁盘总量"

    filesystem_top_group.sort(key=lambda x: x["filesystem_utilization"])
    for i in range(min(topnum, len(filesystem_top_group))):
        row = table_disk_group_asc.add_row()
        row.cells[0].text = filesystem_top_group[i]["groupname"]
        row.cells[1].text = str(filesystem_top_group[i]["hostnum"])
        colored_cell(row.cells[2], filesystem_top_group[i]
                     ["filesystem_utilization"])
        row.cells[3].text = str(convert_unit(
            filesystem_top_group[i]["filesystem_total"]))

    document.add_page_break()
    run_disk_host_desc = document.add_heading("", level=2).add_run(
        "3、磁盘使用率最高的{}台主机".format(topnum))
    run_disk_host_desc.font.name = "微软雅黑"
    run_disk_host_desc._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run_disk_host_desc.font.size = Pt(16)

    # 插入表格 磁盘使用率 主机维度 desc
    table_disk_disk_desc = document.add_table(
        rows=1, cols=4, style="Medium Shading 1 Accent 1")
    table_disk_disk_desc.cell(0, 0).width = Inches(3.0)
    table_disk_disk_desc.cell(0, 1).width = Inches(3.4)
    table_disk_disk_desc.cell(0, 2).width = Inches(1.3)
    table_disk_disk_desc.cell(0, 3).width = Inches(1.3)
    table_disk_disk_desc.cell(0, 0).text = "主机组名称"
    table_disk_disk_desc.cell(0, 1).text = "主机名称"
    table_disk_disk_desc.cell(0, 2).text = "磁盘平均使用率"
    table_disk_disk_desc.cell(0, 3).text = "磁盘总量"

    filesystem_top_host.sort(key=itemgetter("hostname"))
    filesystem_top_host_groupby = []
    for hostname, hosts_iter in groupby(filesystem_top_host, key=itemgetter("hostname")):
        hosts = list(hosts_iter)
        filesystem_top_host_groupby.append({
            "hostname": hostname,
            "filesystem_utilization": hosts[0]["filesystem_utilization"],
            "filesystem_total": hosts[0]["filesystem_total"],
            "groupname": ','.join(h['groupname'] for h in hosts)
        })
    filesystem_top_host_groupby.sort(
        key=itemgetter("filesystem_utilization"), reverse=True)
    for i in range(min(topnum, len(filesystem_top_host_groupby))):
        row = table_disk_disk_desc.add_row()
        row.cells[0].text = filesystem_top_host_groupby[i]["groupname"]
        row.cells[1].text = filesystem_top_host_groupby[i]["hostname"]
        colored_cell(row.cells[2], filesystem_top_host_groupby[i]
                     ["filesystem_utilization"])
        row.cells[3].text = str(convert_unit(
            filesystem_top_host_groupby[i]["filesystem_total"]))

    document.add_page_break()
    run_disk_host_asc = document.add_heading("", level=2).add_run(
        "4、磁盘使用率最低的{}台主机".format(topnum))
    run_disk_host_asc.font.name = "微软雅黑"
    run_disk_host_asc._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run_disk_host_asc.font.size = Pt(16)

    # 插入表格 磁盘使用率 主机维度 asc
    table_disk_disk_asc = document.add_table(
        rows=1, cols=4, style="Medium Shading 1 Accent 1")
    table_disk_disk_asc.cell(0, 0).width = Inches(3.0)
    table_disk_disk_asc.cell(0, 1).width = Inches(3.4)
    table_disk_disk_asc.cell(0, 2).width = Inches(1.3)
    table_disk_disk_asc.cell(0, 3).width = Inches(1.3)
    table_disk_disk_asc.cell(0, 0).text = "主机组名称"
    table_disk_disk_asc.cell(0, 1).text = "主机名称"
    table_disk_disk_asc.cell(0, 2).text = "磁盘平均使用率"
    table_disk_disk_asc.cell(0, 3).text = "磁盘总量"

    filesystem_top_host_groupby.sort(key=itemgetter("filesystem_utilization"))

    for i in range(min(topnum, len(filesystem_top_host_groupby))):
        row = table_disk_disk_asc.add_row()
        row.cells[0].text = filesystem_top_host_groupby[i]["groupname"]
        row.cells[1].text = filesystem_top_host_groupby[i]["hostname"]
        colored_cell(row.cells[2], filesystem_top_host_groupby[i]
                     ["filesystem_utilization"])
        row.cells[3].text = str(convert_unit(
            filesystem_top_host_groupby[i]["filesystem_total"]))

    # cpu使用率低于1%的虚拟机列表
    document.add_page_break()
    run_vm = document.add_heading("", level=1).add_run("七、CPU使用率较低的虚拟机")
    run_vm.font.name = "微软雅黑"
    run_vm.font.size = Pt(20)
    run_vm._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run_vm_cpu = document.add_heading("", level=2).add_run(
        "1、CPU使用率低于1%的虚拟机")
    run_vm_cpu.font.name = "微软雅黑"
    run_vm_cpu._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run_vm_cpu.font.size = Pt(16)

    # 插入表格
    vm_page_num = 1
    vm_page_row_count = 0
    vm_nodata = True

    vm_hosts = api.host.get({
        "with_monitored_items": True,
        "searchInventory": {
            "type": "VM"
        },
        "filter": {
            "available": 1,
            "flag": 0
        },
        "output": ["hostid"]
    })
    for vm in vm_hosts:
        vm_cpu_info = [
            host for host in cpu_top_host if host["hostid"] == vm["hostid"]]
        if not vm_cpu_info:
            continue
        if vm_cpu_info[0].get("cpu_utilization", 0) < 1:
            if vm_page_row_count == 0:
                table_vm_cpu = add_vm_table(
                    document, "Medium Shading 1 Accent 1")
            row = table_vm_cpu.add_row()
            row.cells[0].text = vm_cpu_info[0]["groupname"]
            row.cells[1].text = vm_cpu_info[0]["hostname"]
            colored_cell(row.cells[2], vm_cpu_info[0]["cpu_utilization"])
            vm_page_row_count += 1
            if (vm_page_num == 1 and vm_page_row_count >= 17) or (vm_page_row_count >= 21):
                # 第一页满17行换页 非第一页满21行换页
                vm_page_num += 1
                vm_page_row_count = 0
                vm_nodata = False
                document.add_page_break()
    # 无数据则填充一行`无`
    if vm_nodata:
        table_vm_cpu = add_vm_table(
            document, "Medium Shading 1 Accent 1")
        row = table_vm_cpu.add_row()
        for i in range(len(table_vm_cpu.columns)):
            row.cells[i].text = "无"

    # 设置纸张方向为横向
    for section in document.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

    # 遍历所有表格, 为每个单元格添加边框, 设置文字居中
    for _, table in enumerate(document.tables):
        for r, row in enumerate(table.rows):
            for c, cell in enumerate(row.cells):
                if c == 0:
                    continue  # 跳过第一列
                if r == 0:
                    # 表头用浅色边框
                    color = "#DDDDDD"
                else:
                    # 深色边框
                    color = "#7BA0CD"
                set_cell_border(
                    cell,
                    start={"sz": 1, "color": color, "val": "single"}
                )
                # 除第一列外，表格数据居中显示
                if c > 0:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 调整各级别告警数量表格样式
    for row in range(1, 3):
        for col in range(6):
            table_detail_number.cell(
                row, col).paragraphs[0].runs[0].font.size = Pt(12)
            table_detail_number.cell(
                row, col).paragraphs[0].runs[0].bold = False
            table_detail_number.cell(
                row, col).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table_detail_number.cell(
                row, col).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.save(path)


def getcalc(api, itemids, time_from, time_till):
    """ 计算监控项指定时间范围内的最大值，最小值，平均值

    Args:
        api (ZAPI): zabbix api
        itemids (str | list): itemids
        time_from (str | float): 开始时间戳
        time_till (str | float): 结束时间戳

    Returns:
        tuple: (min, max, avg)
    """
    trends = api.trend.get({
        "itemids": itemids,
        "time_from": time_from,
        "time_till": time_till
    })
    if len(trends) != 0:
        values_min = []
        values_max = []
        values_avg = []
        for trend in trends:
            values_min.append(float(trend["value_min"]))
            values_max.append(float(trend["value_max"]))
            values_avg.append(float(trend["value_avg"]))
        num = len(values_avg)
        avg_value = round(sum(values_avg) / num, 2)
        min_value = min(values_min)
        max_value = max(values_max)
        return min_value, max_value, avg_value
    return 0, 0, 0


def main(args):

    zapi = args.zapi

    output_docx = args.output \
        or "运管平台统计分析月报{}年{}月.docx".format(args.start[:4], args.start[4:6])

    try:
        server_host = zapi.host.get({"filter": {"host": "Zabbix server"}})[0]

        get_word(zapi, server_host["hostid"], output_docx,
                 args.start, args.end, args.topnum)

        logging.info("word 报表导出完成")

        if args.api_secret:
            convertapi.api_secret = args.api_secret
            result = convertapi.convert('pdf', {'File': output_docx})
            print(result)
            # save to pdf
            result.file.save(output_docx.replace('.docx', '.pdf'))
            logging.info("pdf 报表导出完成")

    except Exception:
        logging.error(traceback.format_exc())
        sys.exit(-1)
