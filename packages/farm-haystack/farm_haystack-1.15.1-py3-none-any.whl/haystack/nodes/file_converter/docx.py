from typing import Dict, Optional, List

import logging
from pathlib import Path
import docx

from haystack.nodes.file_converter.base import BaseConverter
from haystack.schema import Document


logger = logging.getLogger(__name__)


class DocxToTextConverter(BaseConverter):
    def convert(
        self,
        file_path: Path,
        meta: Optional[Dict[str, str]] = None,
        remove_numeric_tables: Optional[bool] = None,
        valid_languages: Optional[List[str]] = None,
        encoding: Optional[str] = None,
        id_hash_keys: Optional[List[str]] = None,
    ) -> List[Document]:
        """
        Extract text from a .docx file.
        Note: As docx doesn't contain "page" information, we actually extract and return a list of paragraphs here.
        For compliance with other converters we nevertheless opted for keeping the methods name.

        :param file_path: Path to the .docx file you want to convert
        :param meta: dictionary of meta data key-value pairs to append in the returned document.
        :param remove_numeric_tables: This option uses heuristics to remove numeric rows from the tables.
                                      The tabular structures in documents might be noise for the reader model if it
                                      does not have table parsing capability for finding answers. However, tables
                                      may also have long strings that could possible candidate for searching answers.
                                      The rows containing strings are thus retained in this option.
        :param valid_languages: validate languages from a list of languages specified in the ISO 639-1
                                (https://en.wikipedia.org/wiki/ISO_639-1) format.
                                This option can be used to add test for encoding errors. If the extracted text is
                                not one of the valid languages, then it might likely be encoding error resulting
                                in garbled text.
        :param encoding: Not applicable
        :param id_hash_keys: Generate the document id from a custom list of strings that refer to the document's
            attributes. If you want to ensure you don't have duplicate documents in your DocumentStore but texts are
            not unique, you can modify the metadata and pass e.g. `"meta"` to this field (e.g. [`"content"`, `"meta"`]).
            In this case the id will be generated by using the content and the defined metadata.
        """
        if remove_numeric_tables is None:
            remove_numeric_tables = self.remove_numeric_tables
        if valid_languages is None:
            valid_languages = self.valid_languages
        if remove_numeric_tables is True:
            raise Exception("'remove_numeric_tables' is not supported by DocxToTextConverter.")
        if valid_languages is True:
            raise Exception("Language validation using 'valid_languages' is not supported by DocxToTextConverter.")
        if id_hash_keys is None:
            id_hash_keys = self.id_hash_keys

        file = docx.Document(file_path)  # Creating word reader object.
        paragraphs = [para.text for para in file.paragraphs]
        text = "\n".join(paragraphs)
        document = Document(content=text, meta=meta, id_hash_keys=id_hash_keys)
        return [document]
