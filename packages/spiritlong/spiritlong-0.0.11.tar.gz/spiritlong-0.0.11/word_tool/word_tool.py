## 
# Copyright (c) 2023 Chongqing Spiritlong Technology Co., Ltd.
# All rights reserved.
# 
# @author	arthuryang
# @brief	word工具集，只支持docx。
##

import	docx		# pip install python-docx
import	re

## 打开docx文件
def open_docx(filename):
	try:
		doc	= docx.Document(filename)
		return doc
	except Exception as ex:
		print(str(ex))
		exit		

## 保存docx文件
def save_docx(doc, filename):
	try:
		doc.save(filename)
	except Exception as ex:
		print(f"保存失败：{filename}")
		print(str(ex))
		exit

## 替换docx中两个段落之间的内容
#	match_start	前一段落的特征字符串
#	match_end	后一段落的特征字符串
#	paragraphs	要写入的段落
def insert_paragraphs(doc, match_start, match_end, paragraphs=None):
	p_start_index	= -1
	p_end_index	= -1
	# for i in range(len(doc.paragraphs)):
	# 	if re.findall(match_start, doc.paragraphs[i].text):
	# 		p_start	= i
	# 	if re.findall(match_end, doc.paragraphs[i].text):
	# 		p_end	= i
	
	for p in doc.paragraphs:
		if re.findall(match_start, p.text):
			p_start	= p
		if re.findall(match_end, p.text) and p_start:
			# 确保p_start和p_end的先后顺序
			p_end	= p 

	if p_start and p_end:
		# 都找到了，删除
		flag	= False
		for p in doc.paragraphs:
			if flag and p._element==p_end._element:
				# 结束段落
				flag	= False
			elif p._element==p_start._element:
				# 开始段落
				flag	= True
			elif flag:
				# 该删除
				delete_paragraph(p)
		# 在开始段落之后添加
		add_paragraphs(p_start, paragraphs)

## 删除段落
def delete_paragraph(p):
	# element才是实质的段落数据对象，Paragraph只是指向它的对象封装
	element	= p._element
	element.getparent().remove(element)
	p._p = p._element = None

## 添加多个段落
#	p_start		在此段落之后插入多个段落
#	paragraphs	列表成员每个是一个(runs, style)元组，其中runs列表成员是一个字典
def add_paragraphs(p_start, paragraphs):
	last_p	= p_start
	for p in paragraphs:
		new_p = docx.oxml.xmlchemy.OxmlElement("w:p")
		last_p._p.addnext(new_p)
		last_p	= docx.text.paragraph.Paragraph(new_p, last_p._parent)
		
		if p:
			# 添加此段落的各部分
			for r in p['runs']:
				new_run	= last_p.add_run(r['text'])
				new_run.font.name	= r['font_name']
				new_run.font.size	= r['font_size']
				new_run._element.rPr.rFonts.set(docx.oxml.ns.qn("w:eastAsia"), new_run.font.name)
				new_run.font.color.rgb	= r['font_color_rgb']
				new_run.font.bold	= r['bold']
			# 添加此段落风格
			last_p.alignment				= p['alignment']
			last_p.paragraph_format.left_indent		= p['left_indent']
			last_p.paragraph_format.right_indent		= p['right_indent']
			last_p.paragraph_format.first_line_indent	= p['first_line_indent']
			last_p.paragraph_format.space_before 		= p['space_before']
			last_p.paragraph_format.space_after  		= p['space_after']
			last_p.paragraph_format.line_spacing 		= p['line_spacing']
			last_p.paragraph_format.line_spacing_rule 	= p['line_spacing_rule']
			
			if 'style' in p:
				last_p.style	= p['style']

## 返回默认段落
def new_paragraph_data(runs=[], 
		alignment=docx.enum.text.WD_PARAGRAPH_ALIGNMENT.LEFT, 
		first_line_indent	= 0,
		line_spacing_rule	= docx.enum.text.WD_LINE_SPACING.SINGLE,
		line_spacing		= docx.shared.Pt(10)	):
	return {
		'runs'			: runs,
		'alignment'		: alignment,
		'left_indent'		: docx.shared.Cm(0),
		'right_indent'		: docx.shared.Cm(0),
		'first_line_indent'	: docx.shared.Pt(first_line_indent),
		'space_before'		: docx.shared.Pt(0),
		'space_after'		: docx.shared.Pt(0),
		'line_spacing'		: line_spacing,
		'line_spacing_rule'	: line_spacing_rule,			
	}

## 返回默认run
def new_run_data(text="", font_name="宋体", font_size=12, color=docx.shared.RGBColor(0, 0, 0), bold=False, italic=False):
	return {
		"text"			: text,
		"bold"			: bold,
		"italic"		: italic,	
		"font_name"		: font_name,
		"font_size"		: docx.shared.Pt(font_size),
		"font_color_rgb"	: color,
	}		

## 删除word表格中的行
#	table	表对象
#	line	行序号，从0开始
def table_remove_row(table, line):
	if line>=len(table.rows):
		return 
	row	= table.rows[line]
	row._element.getparent().remove(row._element)

## 设置word表格中的单元格内容，包括字体和大小
def table_cell_fill(table, row, column, content, font_name, font_size):
	cell		= table.cell(row, column)
	# 先清空该表格
	cell.text	= ""

	paragraphs	= [new_paragraph_data([new_run_data(content, font_name=font_name, font_size=font_size)], line_spacing_rule=docx.enum.text.WD_LINE_SPACING.EXACTLY, line_spacing=docx.shared.Pt(20))]
	add_paragraphs(cell.paragraphs[0], paragraphs)
	delete_paragraph(cell.paragraphs[0])

# 调试/测试代码
if __name__ == '__main__':
	pass

